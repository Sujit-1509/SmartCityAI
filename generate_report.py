"""
generate_report.py — Generates JanSevaAI_Project_Report.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)  # Navy
    return h

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A3C6D')
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacing
    return table


# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════

doc.add_paragraph()  # top spacing

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('JanSevaAI')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AI-Powered Municipal Complaint Intelligence Platform')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('Project Report')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Prepared for: ').bold = True
meta.add_run('VISHWANOVA 2026\n')
meta.add_run('National Level Project Competition\n').font.size = Pt(10)
meta.add_run('MIT-WPU, Pune\n\n').font.size = Pt(10)
meta.add_run('Live Demo: ').bold = True
meta.add_run('https://d1lggct31hc8gn.cloudfront.net\n')
meta.add_run('Repository: ').bold = True
meta.add_run('https://github.com/Sujit-1509/jansevaAI')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════

add_heading_styled('Table of Contents', level=1)
toc_items = [
    '1. Introduction',
    '2. Problem Statement',
    '3. Proposed Solution',
    '4. System Architecture',
    '5. Technology Stack',
    '6. Key Features',
    '7. AI Pipeline — How It Works',
    '8. User Roles & Access Control',
    '9. Deployment Infrastructure',
    '10. Impact & Future Scope',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════

add_heading_styled('1. Introduction', level=1)
doc.add_paragraph(
    'JanSevaAI is an end-to-end, AI-powered municipal complaint management platform built '
    'to modernize how Indian cities handle civic grievances. The platform enables citizens to '
    'report infrastructure issues — such as potholes, garbage dumping, broken streetlights, '
    'and waterlogging — by simply uploading a photograph from their smartphone.'
)
doc.add_paragraph(
    'Within seconds, the platform\'s AI pipeline automatically identifies the type of issue, '
    'estimates its severity, generates a formal complaint description, and routes it to the '
    'appropriate municipal department. Administrators can track complaints in real-time, assign '
    'field workers, and monitor resolution progress with built-in SLA tracking. Workers can '
    'resolve complaints by uploading proof-of-resolution photos, which are then verified by AI '
    'to prevent fraudulent closures.'
)
doc.add_paragraph(
    'The name "JanSeva" (जनसेवा) means "Service of the People" in Hindi, reflecting our mission '
    'to bring transparency, speed, and accountability to civic governance through artificial intelligence.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════════════════

add_heading_styled('2. Problem Statement', level=1)
doc.add_paragraph(
    'Urban India generates millions of civic complaints annually, yet traditional grievance redressal '
    'systems suffer from fundamental inefficiencies that leave citizens frustrated and municipalities '
    'overwhelmed:'
)

problems = [
    ('Inaccessible Reporting', 'Citizens must navigate complex government portals, wait in '
     'long phone queues, or visit municipal offices in person to file complaints. This discourages reporting.'),
    ('Manual Triage Bottleneck', 'Every complaint is manually read, categorized, and assigned by office '
     'staff — a process that takes hours or days, causing cascading delays.'),
    ('No Accountability', 'Once filed, complaints enter a black box. Citizens receive no real-time '
     'updates, and there is no mechanism to verify whether reported issues have genuinely been resolved.'),
    ('Fraudulent Closures', 'Field workers can mark tasks as "resolved" without proof, leading to '
     'inflated resolution metrics while physical issues persist on the ground.'),
    ('Zero Analytics', 'Administrators lack data-driven insights into complaint hotspots, recurring '
     'issue patterns, and departmental performance — making strategic planning impossible.'),
]
for title, desc in problems:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}: ')
    run.bold = True
    p.add_run(desc)

# ═══════════════════════════════════════════════════════════════════════════
# 3. PROPOSED SOLUTION
# ═══════════════════════════════════════════════════════════════════════════

add_heading_styled('3. Proposed Solution', level=1)
doc.add_paragraph(
    'JanSevaAI re-imagines the entire complaint lifecycle as an AI-first, serverless pipeline. '
    'Our solution addresses each problem identified above with a targeted technological response:'
)

solutions = [
    ('Photo-Based Submission', 'Citizens simply take a photo of the issue and upload it. The AI handles '
     'classification, severity estimation, and formal complaint generation — zero paperwork required.'),
    ('Automated AI Triage', 'A dual-AI pipeline (YOLOv8 + Amazon Bedrock Nova) classifies complaints '
     'in under 3 seconds, assigning them to the correct municipal department automatically.'),
    ('End-to-End Transparency', 'Every complaint has a complete audit trail — from submission through '
     'assignment, worker acceptance, resolution, and AI-powered verification. Citizens can track '
     'progress in real-time.'),
    ('AI Proof-of-Resolution', 'Workers must upload "after" photos as proof. Our verification AI '
     'compares before/after images to confirm the issue was genuinely fixed at the correct location.'),
    ('Admin Analytics Dashboard', 'Real-time analytics provide category breakdowns, resolution trends, '
     'SLA compliance rates, and geographic hotspot maps for data-driven governance.'),
]
for title, desc in solutions:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}: ')
    run.bold = True
    p.add_run(desc)

# ═══════════════════════════════════════════════════════════════════════════
# 4. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════

add_heading_styled('4. System Architecture', level=1)
doc.add_paragraph(
    'JanSevaAI follows a fully serverless, event-driven architecture deployed entirely on AWS. '
    'The platform consists of three layers:'
)

p = doc.add_paragraph()
p.add_run('Frontend Layer: ').bold = True
p.add_run('A React 19 single-page application built with Vite 7, hosted on Amazon S3 and served '
          'globally through Amazon CloudFront CDN for sub-100ms load times.')

p = doc.add_paragraph()
p.add_run('API Layer: ').bold = True
p.add_run('Amazon API Gateway exposes RESTful endpoints backed by 13 specialized AWS Lambda functions '
          '(Python 3.12). Each Lambda handles a specific domain: authentication, complaint submission, '
          'status updates, worker management, analytics, etc.')

p = doc.add_paragraph()
p.add_run('AI Layer: ').bold = True
p.add_run('A dual-model inference pipeline. The primary path sends images to a custom YOLOv8 model '
          'running on an EC2 instance via FastAPI. If the YOLO server is unavailable or returns low-confidence '
          'results, the system automatically falls back to Amazon Bedrock Nova Lite for vision-based '
          'classification and Amazon Nova Micro for text generation.')

p = doc.add_paragraph()
p.add_run('Data Layer: ').bold = True
p.add_run('Amazon DynamoDB stores all complaint records, user sessions, and worker registrations. '
          'Amazon S3 stores complaint images with server-side encryption.')

# Architecture diagram (text-based)
arch = doc.add_paragraph()
arch.alignment = WD_ALIGN_PARAGRAPH.CENTER
arch_text = (
    'Citizen App → CloudFront → S3 (Frontend)\n'
    '         ↓\n'
    '   API Gateway (REST)\n'
    '         ↓\n'
    '   13 Lambda Functions → DynamoDB\n'
    '         ↓\n'
    '   S3 (Image Upload) → S3 Event Trigger\n'
    '         ↓\n'
    '   Process Image Lambda\n'
    '    ↙              ↘\n'
    'YOLOv8 (EC2)    Bedrock Nova (Fallback)\n'
)
run = arch.add_run(arch_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# ═══════════════════════════════════════════════════════════════════════════
# 5. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════════════════

add_heading_styled('5. Technology Stack', level=1)

add_styled_table(
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Frontend', 'React 19 + Vite 7', 'Single-page application with responsive UI'],
        ['Styling', 'Vanilla CSS Design System', 'Glassmorphic theme with micro-animations'],
        ['API Gateway', 'AWS API Gateway (REST)', 'Unified REST endpoint routing'],
        ['Compute', 'AWS Lambda (Python 3.12)', '13 serverless microservices'],
        ['Database', 'Amazon DynamoDB', 'NoSQL storage for complaints, users, workers'],
        ['Storage', 'Amazon S3', 'Complaint photo storage with encryption'],
        ['CDN', 'Amazon CloudFront', 'Global frontend delivery with edge caching'],
        ['Auth', 'AWS SNS + Custom JWT', 'Phone-based OTP login with HMAC-SHA256 tokens'],
        ['Primary AI', 'YOLOv8 on EC2', 'Custom-trained civic issue detection model'],
        ['Fallback AI', 'Amazon Bedrock Nova Lite', 'Vision classification + spam detection'],
        ['Text AI', 'Amazon Bedrock Nova Micro', 'Formal complaint text generation'],
        ['Notifications', 'Amazon SES + SNS', 'Email and SMS complaint status alerts'],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════
# 6. KEY FEATURES
# ═══════════════════════════════════════════════════════════════════════════

add_heading_styled('6. Key Features', level=1)

add_styled_table(
    ['Feature', 'Description'],
    [
        ['AI Photo Triage', 'Auto-classifies complaint type (pothole, garbage, streetlight, water) from photos'],
        ['Eco-Mode Fallback', 'Seamlessly falls back to Amazon Bedrock when YOLO server is unavailable'],
        ['AI Spam Detection', 'Filters out irrelevant uploads (selfies, memes, screenshots) using Nova Vision'],
        ['Severity Estimation', 'AI-calculated severity scoring (Low/Medium/High/Critical) with priority ranking'],
        ['Complaint Text Generation', 'AI generates formal, location-aware complaint descriptions automatically'],
        ['OTP Authentication', 'Phone-based SMS verification for secure, passwordless login'],
        ['Role-Based Access Control', 'Citizen, Admin, and Worker roles with route guards and API-level enforcement'],
        ['Worker Task Management', 'Accept/reject assignments, resolve with proof photos and GPS coordinates'],
        ['AI Resolution Verification', 'Before/after photo comparison to verify genuine issue resolution'],
        ['Geofence Anti-Fraud', 'GPS location validation ensures worker presence at complaint site'],
        ['SLA Monitoring', 'Automatic warning and breach flagging based on resolution time windows'],
        ['Admin Analytics', 'Complaint trends, category mix, department performance, and hotspot mapping'],
        ['CSV Export', 'Filtered data export for governance reporting and compliance'],
        ['Multi-Photo Support', 'Up to 5 supporting evidence photos per complaint'],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════
# 7. AI PIPELINE
# ═══════════════════════════════════════════════════════════════════════════

add_heading_styled('7. AI Pipeline — How It Works', level=1)
doc.add_paragraph(
    'The AI pipeline is the core intelligence engine of JanSevaAI. It processes every uploaded image '
    'through a multi-stage workflow:'
)

steps = [
    ('Step 1 — Image Upload', 'The citizen uploads a photo. The frontend requests a pre-signed S3 URL '
     'from a Lambda function and uploads the image directly to S3, bypassing the API Gateway payload limit.'),
    ('Step 2 — S3 Event Trigger', 'The S3 bucket fires an event notification when a new object is created. '
     'This triggers the "process_image" Lambda automatically.'),
    ('Step 3 — YOLO Inference (Primary)', 'The Lambda sends the image to a custom-trained YOLOv8 model '
     'running on an EC2 instance via FastAPI. The model classifies the image into one of four categories: '
     'pothole, garbage, streetlight, or water.'),
    ('Step 4 — AI Sanitization & Fallback', 'If YOLO returns low confidence (<75%) or "Unknown", the image '
     'is sent to Amazon Bedrock Nova Lite for a second opinion. Nova also performs spam detection, filtering '
     'out irrelevant images like selfies and screenshots.'),
    ('Step 5 — Severity & Department Mapping', 'Based on the classification and confidence score, the system '
     'calculates a severity level (Low/Medium/High/Critical) and maps the complaint to the appropriate '
     'municipal department (PWD, Sanitation, Electrical, Water Resources).'),
    ('Step 6 — Text Generation', 'Amazon Bedrock Nova Micro generates a formal, professional complaint '
     'description using the category, severity, and location address as context.'),
    ('Step 7 — Storage & Notification', 'The enriched complaint record is stored in DynamoDB. Email and SMS '
     'notifications are dispatched to the citizen via Amazon SES and SNS.'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

# ═══════════════════════════════════════════════════════════════════════════
# 8. USER ROLES
# ═══════════════════════════════════════════════════════════════════════════

add_heading_styled('8. User Roles & Access Control', level=1)
doc.add_paragraph(
    'JanSevaAI implements a three-tier role-based access model enforced at both the frontend (route guards) '
    'and backend (Lambda authorization middleware) levels:'
)

add_styled_table(
    ['Role', 'Capabilities'],
    [
        ['Citizen', 'Submit complaints with photos, track complaint status, view nearby issues, upvote complaints'],
        ['Administrator', 'View all complaints, assign workers, update statuses, manage worker registry, '
         'access analytics dashboard, bulk operations, CSV export'],
        ['Field Worker', 'View assigned tasks, accept/reject assignments, resolve complaints with proof photos '
         'and GPS location, track personal SLA compliance'],
    ]
)

doc.add_paragraph(
    'Authentication is handled through phone-based OTP verification via AWS SNS. Upon successful verification, '
    'a custom HMAC-SHA256 JWT token is issued with a 7-day expiry, containing the user\'s phone number and role. '
    'All subsequent API requests include this token in the Authorization header, and every Lambda independently '
    'verifies the token\'s signature and expiration before processing any request.'
)

# ═══════════════════════════════════════════════════════════════════════════
# 9. DEPLOYMENT
# ═══════════════════════════════════════════════════════════════════════════

add_heading_styled('9. Deployment Infrastructure', level=1)
doc.add_paragraph(
    'JanSevaAI is deployed as a production-ready system on AWS with the following infrastructure:'
)

infra = [
    ('Frontend Hosting', 'React build artifacts are stored in an S3 bucket (smartcityai-frontend-prod) '
     'and served globally via CloudFront CDN with automatic cache invalidation on deployment.'),
    ('API Gateway', 'A single REST API Gateway with path-based routing dispatches requests to 13 Lambda '
     'functions. CORS is configured to allow requests from the CloudFront domain.'),
    ('Lambda Functions', '13 Python 3.12 Lambda functions handle all business logic. Each function is independently '
     'deployable with its own IAM role and environment variables.'),
    ('YOLOv8 Server', 'A t3.medium EC2 instance runs the custom-trained YOLOv8 model via FastAPI. An Elastic IP '
     'ensures a fixed address, and a systemd service guarantees automatic restart on reboot.'),
    ('Database', 'Three DynamoDB tables — Complaints, Users, and Workers — provide low-latency, '
     'serverless data storage with on-demand capacity.'),
]
for title, desc in infra:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}: ')
    run.bold = True
    p.add_run(desc)

p = doc.add_paragraph()
p.add_run('Production URL: ').bold = True
p.add_run('https://d1lggct31hc8gn.cloudfront.net')

# ═══════════════════════════════════════════════════════════════════════════
# 10. IMPACT & FUTURE SCOPE
# ═══════════════════════════════════════════════════════════════════════════

add_heading_styled('10. Impact & Future Scope', level=1)

doc.add_heading('Current Impact', level=2)
impacts = [
    'Reduces complaint triage time from hours to under 3 seconds using AI automation.',
    'Eliminates fraudulent resolution closures through AI-powered before/after photo verification.',
    'Provides citizens with real-time transparency into complaint status and resolution progress.',
    'Gives administrators data-driven insights for proactive urban planning and resource allocation.',
    'Lowers operational costs through a fully serverless architecture with pay-per-use pricing.',
]
for item in impacts:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Future Scope', level=2)
future = [
    'Multilingual support (Hindi, Marathi, Tamil) for complaint submission and notifications.',
    'Mobile app (React Native / Flutter) for easier photo capture and GPS integration.',
    'Predictive analytics using historical data to forecast complaint hotspots before they emerge.',
    'Integration with Smart City Mission dashboards for centralized municipal monitoring.',
    'Voice-based complaint submission for citizens who are not comfortable with text or photo interfaces.',
    'Automated worker dispatch using proximity-based assignment algorithms.',
]
for item in future:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run('— JanSevaAI: Making cities responsive, transparent, and AI-ready. —')
run.italic = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)

# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(__file__), 'JanSevaAI_Project_Report.docx')
doc.save(output_path)
print(f"Report saved to: {output_path}")
